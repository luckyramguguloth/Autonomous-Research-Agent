from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_word_report(markdown_text: str, topic: str):
    """
    Converts the markdown report into a structured Word document.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading(topic, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content processing
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # Heading 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # Normal text
            # Simple bold parsing for **text**
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: # Odd parts are bold
                    run.bold = True
                    
    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
